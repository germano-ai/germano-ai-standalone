import sys
import os
import time
import shutil
import concurrent.futures
import fitz  # PyMuPDF
import docx
import torch
import openpyxl
import xlrd
from xlutils.copy import copy
from transformers import pipeline

INPUT_DIR = "input"
OUTPUT_DIR = "output"
PROCESSED_DIR = "processed"
FAILED_DIR = "failed"
EXTRACTED_DIR = "extracted"
STOP_FLAG = "stop.flag"
PID_FILE = "daemon.pid"

import re
import json

SETTINGS = {
    "enable_ocr": True,
    "enable_faces": True,
    "enable_regex": True,
    "ai_threshold": 0.0
}
try:
    if os.path.exists("config.json"):
        with open("config.json", "r", encoding="utf-8") as f:
            SETTINGS.update(json.load(f))
except Exception:
    pass

def get_regex_spans(text):
    if not SETTINGS.get("enable_regex", True):
        return []
    patterns = [
        # Indirizzi generici
        (r'(?i)\b(?:Via|Viale|Piazza|Corso|Largo|Vicolo|P\.zza|C\.so)\s+[A-Za-z0-9\s\'\.,]+\d{1,5}\b', 'private_address'),
        # CAP + Città
        (r'\b\d{5}\s+[A-Za-z]+\b', 'private_address'),
        # Codice Fiscale
        (r'\b[A-Z]{6}\d{2}[A-Z]\d{2}[A-Z]\d{3}[A-Z]\b', 'cf_iva'),
        # Partita IVA
        (r'\bIT\d{11}\b', 'cf_iva'),
        # Email
        (r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,7}\b', 'private_email'),
        # Telefoni italiani
        (r'\b(?:\+39|0039)?\s?(?:3\d{2}[\s\-\.]?\d{6,7}|0\d{1,3}[\s\-\.]?\d{5,8})\b', 'private_phone')
    ]
    spans = []
    if not isinstance(text, str):
        return spans
        
    for pattern, entity_group in patterns:
        try:
            for match in re.finditer(pattern, text):
                spans.append({
                    "word": match.group(0),
                    "start": match.start(),
                    "end": match.end(),
                    "score": 1.0,
                    "entity": entity_group,
                    "entity_group": entity_group
                })
        except: pass
    return spans

CATEGORY_MAP = {
    "private_person": ("[PERSONA]", "censor_person"),
    "private_address": ("[INDIRIZZO]", "censor_address"),
    "private_email": ("[EMAIL]", "censor_email"),
    "private_phone": ("[TELEFONO]", "censor_phone"),
    "cf_iva": ("[COD.FISCALE/P.IVA]", "censor_cf_iva"),
    "account_number": ("[CONTO/CARTA]", "censor_account"),
    "secret": ("[SEGRETO]", "censor_secret"),
    "private_url": ("[URL]", "censor_url"),
    "private_date": ("[DATA]", "censor_date")
}

def filter_and_format_spans(spans):
    filtered = []
    for span in spans:
        group = span.get("entity_group", span.get("entity", "PII"))
        if group in CATEGORY_MAP:
            label, config_key = CATEGORY_MAP[group]
            if not SETTINGS.get(config_key, True):
                continue
            span["replacement_label"] = label
        else:
            span["replacement_label"] = "[DATI_SENSIBILI]"
        filtered.append(span)
    return filtered

def merge_and_sort_spans(spans, text):
    if not spans: return []
    spans.sort(key=lambda x: x['start'])
    merged = [spans[0].copy()]
    for current in spans[1:]:
        prev = merged[-1]
        if current['start'] <= prev['end']:
            prev['end'] = max(prev['end'], current['end'])
            if isinstance(text, str):
                prev['word'] = text[prev['start']:prev['end']]
        else:
            merged.append(current.copy())
    merged.sort(key=lambda x: x['start'], reverse=True)
    return merged

# --- COPY of all process_* functions from anonymizer.py ---
def process_txt(input_file, output_file, classifier, extracted_data):
    with open(input_file, 'r', encoding='utf-8', errors='ignore') as f:
        text = f.read()
    
    if not text.strip():
        with open(output_file, 'w', encoding='utf-8') as f:
            f.write(text)
        return

    spans = classifier(text) + get_regex_spans(text)
    spans = merge_and_sort_spans(filter_and_format_spans(spans), text)
    
    anonymized_text = text
    for span in spans:
        extracted_data.append({"word": span["word"], "entity": span.get("entity_group", span.get("entity", "PII")), "score": float(span["score"])})
        replacement = span.get('replacement_label', '[DATI_SENSIBILI]')
        anonymized_text = anonymized_text[:span['start']] + replacement + anonymized_text[span['end']:]
        
    with open(output_file, 'w', encoding='utf-8') as f:
        f.write(anonymized_text)

def process_pdf(input_file, output_file, classifier, extracted_data, ocr_reader=None):
    doc = None
    try:
        doc = fitz.open(input_file)
        for page in doc:
            text = page.get_text()
            
            # --- BLOCCO RILEVAMENTO VOLTI (OpenCV) ---
            if SETTINGS["enable_faces"]:
                try:
                    import cv2  # type: ignore
                    import numpy as np  # type: ignore
                    # Cattura un'immagine della pagina a risoluzione standard
                    pix = page.get_pixmap()
                    img = np.frombuffer(pix.samples, dtype=np.uint8).reshape(pix.height, pix.width, pix.n)
                    
                    # Converti in scala di grigi per l'algoritmo Haar Cascade
                    if pix.n >= 4:
                        gray = cv2.cvtColor(img, cv2.COLOR_RGBA2GRAY)
                    elif pix.n == 3:
                        gray = cv2.cvtColor(img, cv2.COLOR_RGB2GRAY)
                    else:
                        gray = img
                        
                    face_cascade = cv2.CascadeClassifier(cv2.data.haarcascades + 'haarcascade_frontalface_default.xml')
                    faces = face_cascade.detectMultiScale(gray, scaleFactor=1.1, minNeighbors=4, minSize=(30, 30))
                    
                    if len(faces) > 0:
                        print(f"[{time.strftime('%H:%M:%S')}] Rilevati {len(faces)} volti nella pagina. Applicazione censura biometrica...", flush=True)
                        for (x, y, w, h) in faces:
                            # Margine di sicurezza del 15% attorno al volto
                            margin_x = int(w * 0.15)
                            margin_y = int(h * 0.15)
                            rect = fitz.Rect(x - margin_x, y - margin_y, x + w + margin_x, y + h + margin_y)
                            
                            annot = page.add_redact_annot(rect, fill=(0,0,0), text="redacted_face", text_color=(0,0,0))
                            annot.set_colors(stroke=None, fill=(0,0,0))
                            annot.update()
                            extracted_data.append({"word": "[VOLTO OSCURATO]", "entity": "BIOMETRIC_FACE", "score": 1.0})
                            
                        # Applica le oscurazioni ai pixel prima di procedere col testo
                        page.apply_redactions(images=2, graphics=0)
                except Exception as e:
                    print(f"Errore nel rilevamento volti: {e}", flush=True)
            # --- FINE BLOCCO VOLTI ---
            
            text = page.get_text()
            
            # --- BLOCCO OCR (SE LA PAGINA E' UNA SCANSIONE O IMMAGINE VUOTA) ---
            if not text.strip():
                if not SETTINGS["enable_ocr"]:
                    continue # Salta elaborazione se l'utente ha disabilitato l'OCR
                    
                if ocr_reader is None:
                    continue # Salta se OCR non è disponibile
                    
                print(f"[{time.strftime('%H:%M:%S')}] Rilevata scansione visiva. Attivazione OCR sulla pagina...", flush=True)
                # Aumenta la risoluzione x2 per un OCR perfetto
                pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))
                img_bytes = pix.tobytes("png")
                
                results = ocr_reader.readtext(img_bytes)
                if not results:
                    continue
                    
                ocr_text = " \n ".join([res[1] for res in results])
                spans_classification = classifier(ocr_text) + get_regex_spans(ocr_text)
                processed_words = set()
                
                for span in spans_classification:
                    word = span['word'].strip()
                    if not word or word in processed_words: continue
                    processed_words.add(word)
                    
                    extracted_data.append({"word": span["word"], "entity": span.get("entity_group", span.get("entity", "PII")), "score": float(span["score"])})
                    word_lower = word.lower()
                    
                    for res in results:
                        bbox, text_seg, conf = res
                        if word_lower in text_seg.lower() or text_seg.lower() in word_lower:
                            # Riporta le coordinate OCR (risoluzione x2) alle coordinate PyMuPDF originali
                            x0 = min(bbox[0][0], bbox[3][0]) / 2.0
                            y0 = min(bbox[0][1], bbox[1][1]) / 2.0
                            x1 = max(bbox[1][0], bbox[2][0]) / 2.0
                            y1 = max(bbox[2][1], bbox[3][1]) / 2.0
                            
                            rect = fitz.Rect(x0, y0, x1, y1)
                            
                            replacement = span.get("replacement_label", "[DATI_SENSIBILI]")
                            
                            annot = page.add_redact_annot(rect, text=replacement, text_color=(1,1,1), fill=(0,0,0), cross_out=False)
                            annot.set_colors(stroke=None, fill=(0,0,0))
                            annot.update()
                            
                # Applica alterando irrimediabilmente l'immagine sottostante
                page.apply_redactions(images=2, graphics=0)
                continue
            # --- FINE BLOCCO OCR ---

            # --- BLOCCO TESTO DIGITALE ---
            page_dict = page.get_text("dict")
            spans_classification = classifier(text) + get_regex_spans(text)
            
            processed_words = set()
            
            # Pre-calc all text spans for faster intersection
            all_spans = []
            for block in page_dict.get("blocks", []):
                if "lines" not in block: continue
                for line in block["lines"]:
                    for s in line["spans"]:
                        all_spans.append((fitz.Rect(s["bbox"]), s["size"]))
            
            for span in spans_classification:
                word = span['word'].strip()
                
                # Save all extracted entities
                extracted_data.append({"word": span["word"], "entity": span.get("entity_group", span.get("entity", "PII")), "score": float(span["score"])})
                
                # Skip if empty or already processed to avoid massive slowdowns (O(N^2) explosion)
                if not word or word in processed_words:
                    continue
                processed_words.add(word)
                
                base_str = span.get("replacement_label", "[DATI_SENSIBILI]")
                replacement = base_str
                
                rects = page.search_for(word)
                for rect in rects:
                    fontsize = rect.height * 0.75
                    
                    # Fast intersection check
                    for span_rect, size in all_spans:
                        if span_rect.intersects(rect):
                            fontsize = size
                            break
                    
                    # Aggiunge annotazione trasparente con fontsize esatto
                    annot = page.add_redact_annot(rect, text=replacement, text_color=(0,0,0), fontsize=fontsize, fill=False, cross_out=False)
                    # Forza la rimozione di bordi rossi o sfondi di default
                    annot.set_colors(stroke=None, fill=None)
                    annot.update()
            
            # Applica preservando le immagini intatte
            page.apply_redactions(images=0, graphics=0)
            
        doc.save(output_file, garbage=4, deflate=True, clean=True)
    finally:
        if doc:
            doc.close()

def process_docx(input_file, output_file, classifier, extracted_data):
    doc = docx.Document(input_file)
    
    def process_paragraphs(paragraphs):
        for para in paragraphs:
            if not para.text.strip():
                continue
                
            text = para.text
            spans = classifier(text) + get_regex_spans(text)
            if not spans:
                continue
                
            spans = merge_and_sort_spans(filter_and_format_spans(spans), text)
            anonymized_text = text
            for span in spans:
                extracted_data.append({"word": span["word"], "entity": span.get("entity_group", span.get("entity", "PII")), "score": float(span["score"])})
                replacement = span.get('replacement_label', '[DATI_SENSIBILI]')
                anonymized_text = anonymized_text[:span['start']] + replacement + anonymized_text[span['end']:]
                
            para.text = anonymized_text

    def process_tables(tables):
        for table in tables:
            for row in table.rows:
                for cell in row.cells:
                    process_paragraphs(cell.paragraphs)

    # Elabora il corpo principale del documento
    process_paragraphs(doc.paragraphs)
    process_tables(doc.tables)
    
    # Elabora intestazioni e piè di pagina in tutte le sezioni
    for section in doc.sections:
        headers_footers = [
            section.header, section.first_page_header, section.even_page_header,
            section.footer, section.first_page_footer, section.even_page_footer
        ]
        for hf in headers_footers:
            if hf is not None and not hf.is_linked_to_previous:
                process_paragraphs(hf.paragraphs)
                process_tables(hf.tables)


    doc.save(output_file)

def process_xlsx(input_file, output_file, classifier, extracted_data):
    import os
    from openpyxl.comments import Comment
    
    ext = os.path.splitext(input_file)[1].lower()
    wb = None
    try:
        wb = openpyxl.load_workbook(input_file, keep_vba=(ext == '.xlsm'))
        
        def anonymize_text(text):
            if not isinstance(text, str) or not text.strip():
                return text
            spans = classifier(text) + get_regex_spans(text)
            if not spans:
                return text
            spans = merge_and_sort_spans(filter_and_format_spans(spans), text)
            anonymized_text = text
            for span in spans:
                extracted_data.append({"word": span["word"], "entity": span.get("entity_group", span.get("entity", "PII")), "score": float(span["score"])})
                replacement = span.get('replacement_label', '[DATI_SENSIBILI]')
                anonymized_text = anonymized_text[:span['start']] + replacement + anonymized_text[span['end']:]
            return anonymized_text

        for sheet in wb.worksheets:
            new_title = anonymize_text(sheet.title)
            if new_title and new_title != sheet.title:
                sheet.title = new_title[:31] # Excel limita i nomi a 31 caratteri
                
            for row in sheet.iter_rows():
                for cell in row:
                    if isinstance(cell.value, str):
                        cell.value = anonymize_text(cell.value)
                    
                    if cell.comment and cell.comment.text:
                        new_comment_text = anonymize_text(cell.comment.text)
                        if new_comment_text != cell.comment.text:
                            author = anonymize_text(cell.comment.author) if cell.comment.author else "Author"
                            cell.comment = Comment(new_comment_text, author)

        wb.save(output_file)
    finally:
        if wb and hasattr(wb, 'close'):
            wb.close()

def process_xls(input_file, output_file, classifier, extracted_data):
    try:
        rb = xlrd.open_workbook(input_file, formatting_info=True)
    except NotImplementedError:
        rb = xlrd.open_workbook(input_file)
        
    wb = copy(rb)
    for s in range(rb.nsheets):
        sheet = rb.sheet_by_index(s)
        w_sheet = wb.get_sheet(s)
        
        for rowx in range(sheet.nrows):
            for colx in range(sheet.ncols):
                val = sheet.cell_value(rowx, colx)
                if isinstance(val, str) and val.strip():
                    spans = classifier(val) + get_regex_spans(val)
                    if not spans:
                        continue
                        
                    spans = merge_and_sort_spans(filter_and_format_spans(spans), val)
                    anonymized_text = val
                    for span in spans:
                        extracted_data.append({"word": span["word"], "entity": span.get("entity_group", span.get("entity", "PII")), "score": float(span["score"])})
                        replacement = span.get("replacement_label", "[DATI_SENSIBILI]")
                        anonymized_text = anonymized_text[:span['start']] + replacement + anonymized_text[span['end']:]
                    
                    w_sheet.write(rowx, colx, anonymized_text)
                    
    wb.save(output_file)

def process_rtf(input_file, output_file, classifier, extracted_data):
    with open(input_file, 'r', encoding='utf-8', errors='ignore') as f:
        text = f.read()
    
    if not text.strip():
        with open(output_file, 'w', encoding='utf-8') as f:
            f.write(text)
        return

    spans = classifier(text) + get_regex_spans(text)
    spans = merge_and_sort_spans(filter_and_format_spans(spans), text)
    
    anonymized_text = text
    for span in spans:
        extracted_data.append({"word": span["word"], "entity": span.get("entity_group", span.get("entity", "PII")), "score": float(span["score"])})
        replacement = span.get('replacement_label', '[DATI_SENSIBILI]')
        anonymized_text = anonymized_text[:span['start']] + replacement + anonymized_text[span['end']:]
        
    with open(output_file, 'w', encoding='utf-8') as f:
        f.write(anonymized_text)


def main():
    # Write PID for PHP to monitor
    with open(PID_FILE, "w") as f:
        f.write(str(os.getpid()))
        
    # Clear old stop flag if exists
    if os.path.exists(STOP_FLAG):
        os.remove(STOP_FLAG)

    print("Initializing Daemon...", flush=True)
    os.makedirs(INPUT_DIR, exist_ok=True)
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    os.makedirs(PROCESSED_DIR, exist_ok=True)
    os.makedirs(FAILED_DIR, exist_ok=True)
    os.makedirs(EXTRACTED_DIR, exist_ok=True)
    
    READY_FLAG = "ready.flag"
    if os.path.exists(READY_FLAG):
        os.remove(READY_FLAG)
    
    print("Loading AI model (germano.ai Core)...", flush=True)
    
    hardware_mode = SETTINGS.get("hardware_mode", "gpu_fallback")
    
    device = "cpu"
    
    if hardware_mode in ["gpu", "gpu_fallback"]:
        # 1. Check for NVIDIA CUDA
        if torch.cuda.is_available():
            device = "cuda"
            print("Hardware Rilevato: NVIDIA GPU (CUDA disponibile). Accelerazione attivata.", flush=True)
        else:
            # 2. Check for AMD / Intel via DirectML
            try:
                import torch_directml  # type: ignore
                if torch_directml.is_available():
                    device = torch_directml.device()
                    print(f"Hardware Rilevato: AMD/Intel GPU (DirectML disponibile). Accelerazione attivata.", flush=True)
                else:
                    print("Nessuna GPU compatibile rilevata. Avvio in modalita' Failsafe su CPU.", flush=True)
            except ImportError:
                print("Librerie grafiche non trovate. Avvio in modalita' Failsafe su CPU.", flush=True)
    else:
        print("Forzatura manuale elaborazione su Processore Centrale (CPU).", flush=True)

    try:
        model_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "models", "privacy-filter")
        
        # Fallback to online if local missing (useful for testing)
        if not os.path.exists(model_path) or not os.listdir(model_path):
            print("Local model not found, falling back to HuggingFace Hub...", flush=True)
            model_path = "openai/privacy-filter"
            
        raw_classifier = pipeline(
            task="token-classification", 
            model=model_path, 
            aggregation_strategy="simple",
            trust_remote_code=True,
            device=device,
            model_kwargs={"dtype": torch.float32}
        )
        
        # Wrapper per intercettare gli errori critici di DirectML su GPU AMD e scalare a CPU
        class SafeClassifier:
            def __init__(self, pipe, threshold=0.0):
                self.pipe = pipe
                self.fallback_done = False
                self.threshold = threshold
                
            def __call__(self, text):
                res = []
                try:
                    res = self.pipe(text)
                except Exception as e:
                    if SETTINGS.get("hardware_mode", "gpu_fallback") == "gpu_fallback" and "DirectML" in str(e) and not self.fallback_done:
                        print(f"[{time.strftime('%H:%M:%S')}] ERRORE GRAVE: Incompatibilita' hardware rilevata. I driver DirectML della GPU hanno fallito ('{str(e)[:50]}...').", flush=True)
                        print(f"[{time.strftime('%H:%M:%S')}] DISATTIVAZIONE GPU IN CORSO. Fallback del modello matematico su Processore Centrale (CPU)...", flush=True)
                        self.pipe.model = self.pipe.model.to("cpu")
                        self.pipe.device = torch.device("cpu")
                        self.fallback_done = True
                        print(f"[{time.strftime('%H:%M:%S')}] Fallback riuscito. Riavvio elaborazione su CPU.", flush=True)
                        res = self.pipe(text)
                    else:
                        raise e
                
                # Applica filtro threshold
                return [s for s in res if s.get("score", 0.0) >= self.threshold]

        classifier = SafeClassifier(raw_classifier, threshold=SETTINGS.get("ai_threshold", 0.0))
        
    except Exception as e:
        print(f"Error loading model: {str(e)}", flush=True)
        sys.exit(1)
        
    print("Loading OCR Engine (EasyOCR)...", flush=True)
    try:
        import easyocr  # type: ignore
        os.environ["EASYOCR_MODULE_PATH"] = os.path.join(os.path.dirname(os.path.abspath(__file__)), "models", "easyocr")
        os.makedirs(os.environ["EASYOCR_MODULE_PATH"], exist_ok=True)
        # Usa GPU solo se c'è CUDA puro. Su AMD (DirectML), meglio CPU per stabilita' dell'OCR.
        use_gpu = True if torch.cuda.is_available() else False
        ocr_reader = easyocr.Reader(['it', 'en'], gpu=use_gpu, verbose=False, download_enabled=False)
    except Exception as e:
        print(f"Error loading OCR Engine: {str(e)}", flush=True)
        ocr_reader = None
        
    print("Model loaded successfully. Entering monitoring loop...", flush=True)
    
    with open(READY_FLAG, "w") as f:
        f.write("1")
        
    print("Ready for processing.", flush=True)
    
    last_activity_time = time.time()
    
    while True:
        if os.path.exists(STOP_FLAG):
            print("Stop flag detected. Shutting down daemon safely...", flush=True)
            os.remove(STOP_FLAG)
            if os.path.exists(PID_FILE):
                os.remove(PID_FILE)
            break
            
        found_files = []
        for root, dirs, files in os.walk(INPUT_DIR):
            for file in files:
                found_files.append(os.path.join(root, file))
                
        if not found_files:
            if time.time() - last_activity_time > 60:
                print("No input for 60 seconds. Auto-shutting down to save resources...", flush=True)
                if os.path.exists(PID_FILE):
                    os.remove(PID_FILE)
                break
            time.sleep(2)
            continue
            
        for input_file in found_files:
            if os.path.exists(STOP_FLAG):
                break
                
            rel_path = os.path.relpath(input_file, INPUT_DIR)
            output_file = os.path.join(OUTPUT_DIR, rel_path)
            processed_file = os.path.join(PROCESSED_DIR, rel_path)
            
            try:
                os.makedirs(os.path.dirname(output_file), exist_ok=True)
                os.makedirs(os.path.dirname(processed_file), exist_ok=True)
                
                print(f"Processing: {rel_path}...", flush=True)
                
                ext = os.path.splitext(input_file)[1].lower()
                success = False
                is_failed = False
                extracted_data = []
                
                def run_task():
                    if ext == '.txt':
                        process_txt(input_file, output_file, classifier, extracted_data)
                    elif ext in ['.pdf', '.png', '.jpg', '.jpeg', '.tiff', '.bmp']:
                        out_pdf = output_file
                        if ext != '.pdf':
                            out_pdf = os.path.splitext(output_file)[0] + ".pdf"
                        process_pdf(input_file, out_pdf, classifier, extracted_data, ocr_reader=ocr_reader)
                    elif ext in ['.docx', '.docm']:
                        process_docx(input_file, output_file, classifier, extracted_data)
                    elif ext in ['.xlsx', '.xlsm']:
                        process_xlsx(input_file, output_file, classifier, extracted_data)
                    elif ext == '.xls':
                        process_xls(input_file, output_file, classifier, extracted_data)
                    elif ext == '.rtf':
                        process_rtf(input_file, output_file, classifier, extracted_data)
                    else:
                        return "unsupported"
                    return "success"

                executor = concurrent.futures.ThreadPoolExecutor(max_workers=1)
                future = executor.submit(run_task)
                try:
                    result = future.result(timeout=300)
                    if result == "unsupported":
                        print(f"Unsupported file type skipped: {rel_path}", flush=True)
                        success = True
                    else:
                        print(f"Success: {rel_path}. Saved to output.", flush=True)
                        success = True
                except concurrent.futures.TimeoutError:
                    print(f"[TIMEOUT ERROR] Il motore si e' bloccato su {rel_path} (piu' di 5 minuti). Spostamento in failed e reset.", flush=True)
                    is_failed = True
                finally:
                    executor.shutdown(wait=False)
                        
                if success:
                    try:
                        import json
                        extracted_file = os.path.join(EXTRACTED_DIR, rel_path + ".json")
                        os.makedirs(os.path.dirname(extracted_file), exist_ok=True)
                        with open(extracted_file, "w", encoding="utf-8") as f:
                            json.dump(extracted_data, f, indent=4, ensure_ascii=False)
                    except Exception as e:
                        print(f"Could not save extracted data for {rel_path}: {str(e)}", flush=True)

                    try:
                        if os.path.exists(processed_file):
                            os.remove(processed_file)
                        shutil.move(input_file, processed_file)
                    except Exception as e:
                        print(f"Could not move {rel_path} to processed: {str(e)}", flush=True)
                elif is_failed:
                    raise Exception("Process failed or timed out.")
                    
            except Exception as e:
                print(f"[CRITICAL ERROR] Impossibile processare {rel_path}: {str(e)}", flush=True)
                try:
                    # Se l'errore è dovuto alla lunghezza del path, spostiamolo con un nome corto
                    failed_file = os.path.join(FAILED_DIR, rel_path)
                    try:
                        os.makedirs(os.path.dirname(failed_file), exist_ok=True)
                    except:
                        import hashlib
                        short_hash = hashlib.md5(rel_path.encode('utf-8')).hexdigest()
                        ext = os.path.splitext(input_file)[1]
                        failed_file = os.path.join(FAILED_DIR, f"error_{short_hash[:8]}{ext}")
                        
                    if os.path.exists(failed_file):
                        os.remove(failed_file)
                    shutil.move(input_file, failed_file)
                except Exception as fatal_e:
                    print(f"FATALE: Impossibile perfino spostare in failed: {str(fatal_e)}", flush=True)

            # Try to clean up empty directories in input
            try:
                input_sub_dir = os.path.dirname(input_file)
                if input_sub_dir != INPUT_DIR and not os.listdir(input_sub_dir):
                    os.rmdir(input_sub_dir)
            except Exception:
                pass
        
        # Reset timeout timer after processing files
        last_activity_time = time.time()

if __name__ == "__main__":
    main()
